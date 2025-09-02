"""
Meeting Export Module
Handles export of meeting data to various formats (PDF, Word, HTML)
"""

import json
from datetime import datetime
from io import BytesIO
from typing import Dict, List, Any
import html

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False


class MeetingExporter:
    """Handles exporting meeting data to various formats"""
    
    def __init__(self):
        self.styles = None
        if REPORTLAB_AVAILABLE:
            self.styles = getSampleStyleSheet()
            
    def format_datetime(self, dt_string: str) -> str:
        """Format datetime string for display"""
        try:
            if not dt_string:
                return "Not specified"
            dt = datetime.fromisoformat(dt_string.replace('Z', '+00:00'))
            return dt.strftime("%B %d, %Y at %I:%M %p")
        except:
            return dt_string or "Not specified"
    
    def format_date(self, date_string: str) -> str:
        """Format date string for display"""
        try:
            if not date_string:
                return "Not specified"
            # Handle various date formats
            if 'T' in date_string:
                dt = datetime.fromisoformat(date_string.replace('Z', '+00:00'))
                return dt.strftime("%B %d, %Y")
            else:
                dt = datetime.strptime(date_string, "%Y-%m-%d")
                return dt.strftime("%B %d, %Y")
        except:
            return date_string or "Not specified"
    
    def clean_html_content(self, content: str) -> str:
        """Clean HTML content for text-based exports"""
        if not content:
            return ""
        
        # Remove HTML tags but preserve some formatting
        import re
        # Replace <br> and <p> tags with newlines
        content = re.sub(r'<br\s*/?>', '\n', content)
        content = re.sub(r'</?p[^>]*>', '\n', content)
        content = re.sub(r'<li[^>]*>', '• ', content)
        content = re.sub(r'</li>', '\n', content)
        # Remove all other HTML tags
        content = re.sub(r'<[^>]+>', '', content)
        # Clean up whitespace
        content = re.sub(r'\n\s*\n', '\n\n', content)
        return content.strip()
    
    def export_to_html(self, meeting: Dict[str, Any]) -> str:
        """Export meeting to formatted HTML"""
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>Meeting Minutes: {html.escape(meeting.get('title', 'Untitled Meeting'))}</title>
            <style>
                body {{
                    font-family: Arial, sans-serif;
                    line-height: 1.6;
                    max-width: 800px;
                    margin: 0 auto;
                    padding: 20px;
                    color: #333;
                }}
                .header {{
                    text-align: center;
                    border-bottom: 2px solid #333;
                    padding-bottom: 20px;
                    margin-bottom: 30px;
                }}
                .meeting-title {{
                    font-size: 24px;
                    font-weight: bold;
                    margin-bottom: 10px;
                }}
                .meeting-info {{
                    font-size: 14px;
                    color: #666;
                }}
                .section {{
                    margin-bottom: 30px;
                }}
                .section-title {{
                    font-size: 18px;
                    font-weight: bold;
                    color: #2c3e50;
                    margin-bottom: 15px;
                    border-bottom: 1px solid #ddd;
                    padding-bottom: 5px;
                }}
                .attendee-list {{
                    display: grid;
                    grid-template-columns: repeat(auto-fill, minmax(200px, 1fr));
                    gap: 10px;
                    margin-bottom: 20px;
                }}
                .attendee {{
                    padding: 8px 12px;
                    background-color: #f8f9fa;
                    border-radius: 4px;
                }}
                .agenda-item {{
                    margin-bottom: 20px;
                    padding: 15px;
                    background-color: #f8f9fa;
                    border-radius: 5px;
                }}
                .agenda-title {{
                    font-weight: bold;
                    margin-bottom: 10px;
                }}
                .action-items-table {{
                    width: 100%;
                    border-collapse: collapse;
                    margin-top: 15px;
                }}
                .action-items-table th,
                .action-items-table td {{
                    border: 1px solid #ddd;
                    padding: 8px;
                    text-align: left;
                }}
                .action-items-table th {{
                    background-color: #f2f2f2;
                    font-weight: bold;
                }}
                .decision {{
                    margin-bottom: 15px;
                    padding: 12px;
                    background-color: #fff3cd;
                    border-left: 4px solid #ffc107;
                    border-radius: 4px;
                }}
                .notes {{
                    background-color: #f8f9fa;
                    padding: 15px;
                    border-radius: 5px;
                    font-style: italic;
                }}
                .print-only {{
                    display: none;
                }}
                @media print {{
                    .print-only {{
                        display: block;
                    }}
                    body {{
                        margin: 0;
                        padding: 15px;
                    }}
                }}
            </style>
        </head>
        <body>
            <div class="header">
                <div class="meeting-title">{html.escape(meeting.get('title', 'Untitled Meeting'))}</div>
                <div class="meeting-info">
                    <strong>Date:</strong> {self.format_datetime(meeting.get('date', ''))} | 
                    <strong>Type:</strong> {html.escape(meeting.get('type', 'Meeting'))} | 
                    <strong>Duration:</strong> {meeting.get('duration', 'N/A')} minutes
                </div>
                {f'<div class="meeting-info"><strong>Location:</strong> {html.escape(meeting.get("location", ""))}</div>' if meeting.get('location') else ''}
                {f'<div class="meeting-info"><strong>Organizer:</strong> {html.escape(meeting.get("organizer", ""))}</div>' if meeting.get('organizer') else ''}
            </div>
        """
        
        # Attendees Section
        if meeting.get('attendees'):
            html_content += """
            <div class="section">
                <div class="section-title">Attendees</div>
                <div class="attendee-list">
            """
            for attendee in meeting['attendees']:
                name = html.escape(attendee.get('name', 'Unknown'))
                role = html.escape(attendee.get('role', ''))
                email = html.escape(attendee.get('email', ''))
                status = html.escape(attendee.get('status', 'Present'))
                
                attendee_info = f"<div class='attendee'><strong>{name}</strong>"
                if role:
                    attendee_info += f"<br><small>{role}</small>"
                if email:
                    attendee_info += f"<br><small>{email}</small>"
                attendee_info += f"<br><small>Status: {status}</small></div>"
                html_content += attendee_info
            
            html_content += "</div></div>"
        
        # Agenda Section
        if meeting.get('agenda'):
            html_content += """
            <div class="section">
                <div class="section-title">Agenda & Discussions</div>
            """
            for i, item in enumerate(meeting['agenda'], 1):
                item_title = html.escape(item.get('item', f'Agenda Item {i}'))
                duration = item.get('duration', '')
                discussion = item.get('discussion', '')
                
                html_content += f"""
                <div class="agenda-item">
                    <div class="agenda-title">{i}. {item_title}"""
                if duration:
                    html_content += f" ({duration} min)"
                html_content += "</div>"
                
                if discussion:
                    html_content += f"<div>{discussion}</div>"
                
                html_content += "</div>"
            
            html_content += "</div>"
        
        # Action Items Section
        if meeting.get('action_items'):
            html_content += """
            <div class="section">
                <div class="section-title">Action Items</div>
                <table class="action-items-table">
                    <thead>
                        <tr>
                            <th>Action Item</th>
                            <th>Assigned To</th>
                            <th>Due Date</th>
                            <th>Status</th>
                        </tr>
                    </thead>
                    <tbody>
            """
            
            for action in meeting['action_items']:
                item_text = html.escape(action.get('item', 'N/A'))
                assignee = html.escape(action.get('assignee', 'Unassigned'))
                due_date = self.format_date(action.get('due_date', ''))
                status = html.escape(action.get('status', 'Open'))
                
                html_content += f"""
                <tr>
                    <td>{item_text}</td>
                    <td>{assignee}</td>
                    <td>{due_date}</td>
                    <td>{status}</td>
                </tr>
                """
            
            html_content += """
                    </tbody>
                </table>
            </div>
            """
        
        # Decisions Section
        if meeting.get('decisions'):
            html_content += """
            <div class="section">
                <div class="section-title">Decisions Made</div>
            """
            
            for i, decision in enumerate(meeting['decisions'], 1):
                decision_text = html.escape(decision.get('decision', ''))
                impact = html.escape(decision.get('impact', ''))
                
                html_content += f"""
                <div class="decision">
                    <strong>Decision {i}:</strong> {decision_text}
                    {f'<br><strong>Impact:</strong> {impact}' if impact else ''}
                </div>
                """
            
            html_content += "</div>"
        
        # Notes Section
        if meeting.get('notes'):
            notes_content = meeting['notes']
            html_content += f"""
            <div class="section">
                <div class="section-title">Additional Notes</div>
                <div class="notes">{notes_content}</div>
            </div>
            """
        
        # Next Meeting Section
        next_meeting = meeting.get('next_meeting', {})
        if next_meeting:
            html_content += f"""
            <div class="section">
                <div class="section-title">Next Meeting</div>
                <p><strong>Date:</strong> {self.format_datetime(next_meeting.get('date', ''))}</p>
                {f'<p><strong>Location:</strong> {html.escape(next_meeting.get("location", ""))}</p>' if next_meeting.get('location') else ''}
                {f'<p><strong>Agenda Items:</strong> {html.escape(next_meeting.get("agenda_preview", ""))}</p>' if next_meeting.get('agenda_preview') else ''}
            </div>
            """
        
        html_content += """
            <div class="print-only">
                <hr>
                <p><small>Generated on """ + datetime.now().strftime("%B %d, %Y at %I:%M %p") + """</small></p>
            </div>
        </body>
        </html>
        """
        
        return html_content
    
    def export_to_pdf(self, meeting: Dict[str, Any]) -> BytesIO:
        """Export meeting to PDF format"""
        if not REPORTLAB_AVAILABLE:
            raise ImportError("reportlab is required for PDF export")
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=1*inch, bottomMargin=1*inch)
        story = []
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=18,
            spaceAfter=20,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=self.styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            textColor=colors.HexColor('#2c3e50')
        )
        
        normal_style = self.styles['Normal']
        
        # Title
        story.append(Paragraph(f"Meeting Minutes: {meeting.get('title', 'Untitled Meeting')}", title_style))
        
        # Meeting Info
        meeting_info = [
            f"<b>Date:</b> {self.format_datetime(meeting.get('date', ''))}",
            f"<b>Type:</b> {meeting.get('type', 'Meeting')}",
            f"<b>Duration:</b> {meeting.get('duration', 'N/A')} minutes"
        ]
        
        if meeting.get('location'):
            meeting_info.append(f"<b>Location:</b> {meeting.get('location')}")
        if meeting.get('organizer'):
            meeting_info.append(f"<b>Organizer:</b> {meeting.get('organizer')}")
        
        for info in meeting_info:
            story.append(Paragraph(info, normal_style))
        
        story.append(Spacer(1, 20))
        
        # Attendees
        if meeting.get('attendees'):
            story.append(Paragraph("Attendees", heading_style))
            
            attendee_data = [['Name', 'Role', 'Email', 'Status']]
            for attendee in meeting['attendees']:
                attendee_data.append([
                    attendee.get('name', 'Unknown'),
                    attendee.get('role', ''),
                    attendee.get('email', ''),
                    attendee.get('status', 'Present')
                ])
            
            attendee_table = Table(attendee_data, colWidths=[2*inch, 1.5*inch, 2*inch, 1*inch])
            attendee_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
            ]))
            story.append(attendee_table)
            story.append(Spacer(1, 20))
        
        # Agenda & Discussions
        if meeting.get('agenda'):
            story.append(Paragraph("Agenda & Discussions", heading_style))
            
            for i, item in enumerate(meeting['agenda'], 1):
                item_title = f"{i}. {item.get('item', f'Agenda Item {i}')}"
                if item.get('duration'):
                    item_title += f" ({item['duration']} min)"
                
                story.append(Paragraph(item_title, ParagraphStyle(
                    'AgendaItem',
                    parent=self.styles['Normal'],
                    fontName='Helvetica-Bold',
                    fontSize=11,
                    spaceAfter=6
                )))
                
                if item.get('discussion'):
                    discussion_clean = self.clean_html_content(item['discussion'])
                    story.append(Paragraph(discussion_clean, normal_style))
                
                story.append(Spacer(1, 10))
        
        # Action Items
        if meeting.get('action_items'):
            story.append(Paragraph("Action Items", heading_style))
            
            action_data = [['Action Item', 'Assigned To', 'Due Date', 'Status']]
            for action in meeting['action_items']:
                action_data.append([
                    action.get('item', 'N/A'),
                    action.get('assignee', 'Unassigned'),
                    self.format_date(action.get('due_date', '')),
                    action.get('status', 'Open')
                ])
            
            action_table = Table(action_data, colWidths=[3*inch, 1.5*inch, 1.5*inch, 1*inch])
            action_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ]))
            story.append(action_table)
            story.append(Spacer(1, 20))
        
        # Decisions
        if meeting.get('decisions'):
            story.append(Paragraph("Decisions Made", heading_style))
            
            for i, decision in enumerate(meeting['decisions'], 1):
                decision_text = f"<b>Decision {i}:</b> {decision.get('decision', '')}"
                if decision.get('impact'):
                    decision_text += f"<br/><b>Impact:</b> {decision.get('impact')}"
                
                story.append(Paragraph(decision_text, normal_style))
                story.append(Spacer(1, 10))
        
        # Notes
        if meeting.get('notes'):
            story.append(Paragraph("Additional Notes", heading_style))
            notes_clean = self.clean_html_content(meeting['notes'])
            story.append(Paragraph(notes_clean, normal_style))
            story.append(Spacer(1, 20))
        
        # Next Meeting
        next_meeting = meeting.get('next_meeting', {})
        if next_meeting:
            story.append(Paragraph("Next Meeting", heading_style))
            next_info = [
                f"<b>Date:</b> {self.format_datetime(next_meeting.get('date', ''))}"
            ]
            if next_meeting.get('location'):
                next_info.append(f"<b>Location:</b> {next_meeting.get('location')}")
            if next_meeting.get('agenda_preview'):
                next_info.append(f"<b>Agenda Items:</b> {next_meeting.get('agenda_preview')}")
            
            for info in next_info:
                story.append(Paragraph(info, normal_style))
        
        # Footer
        story.append(Spacer(1, 30))
        story.append(Paragraph(
            f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}",
            ParagraphStyle(
                'Footer',
                parent=self.styles['Normal'],
                fontSize=8,
                textColor=colors.grey,
                alignment=TA_CENTER
            )
        ))
        
        doc.build(story)
        buffer.seek(0)
        return buffer
    
    def export_to_word(self, meeting: Dict[str, Any]) -> BytesIO:
        """Export meeting to Word document format"""
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError("python-docx is required for Word export")
        
        doc = Document()
        
        # Set up styles
        styles = doc.styles
        
        # Title
        title = doc.add_heading(f"Meeting Minutes: {meeting.get('title', 'Untitled Meeting')}", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Meeting Info
        info_paragraph = doc.add_paragraph()
        info_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        meeting_info = [
            f"Date: {self.format_datetime(meeting.get('date', ''))}",
            f"Type: {meeting.get('type', 'Meeting')}",
            f"Duration: {meeting.get('duration', 'N/A')} minutes"
        ]
        
        if meeting.get('location'):
            meeting_info.append(f"Location: {meeting.get('location')}")
        if meeting.get('organizer'):
            meeting_info.append(f"Organizer: {meeting.get('organizer')}")
        
        info_paragraph.add_run(" | ".join(meeting_info))
        
        doc.add_paragraph()  # Add spacing
        
        # Attendees
        if meeting.get('attendees'):
            doc.add_heading('Attendees', level=1)
            
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Name'
            hdr_cells[1].text = 'Role'
            hdr_cells[2].text = 'Email'
            hdr_cells[3].text = 'Status'
            
            # Make header bold
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Add attendee rows
            for attendee in meeting['attendees']:
                row_cells = table.add_row().cells
                row_cells[0].text = attendee.get('name', 'Unknown')
                row_cells[1].text = attendee.get('role', '')
                row_cells[2].text = attendee.get('email', '')
                row_cells[3].text = attendee.get('status', 'Present')
            
            doc.add_paragraph()  # Add spacing
        
        # Agenda & Discussions
        if meeting.get('agenda'):
            doc.add_heading('Agenda & Discussions', level=1)
            
            for i, item in enumerate(meeting['agenda'], 1):
                item_title = f"{i}. {item.get('item', f'Agenda Item {i}')}"
                if item.get('duration'):
                    item_title += f" ({item['duration']} min)"
                
                agenda_heading = doc.add_heading(item_title, level=2)
                
                if item.get('discussion'):
                    discussion_clean = self.clean_html_content(item['discussion'])
                    doc.add_paragraph(discussion_clean)
        
        # Action Items
        if meeting.get('action_items'):
            doc.add_heading('Action Items', level=1)
            
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Action Item'
            hdr_cells[1].text = 'Assigned To'
            hdr_cells[2].text = 'Due Date'
            hdr_cells[3].text = 'Status'
            
            # Make header bold
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Add action item rows
            for action in meeting['action_items']:
                row_cells = table.add_row().cells
                row_cells[0].text = action.get('item', 'N/A')
                row_cells[1].text = action.get('assignee', 'Unassigned')
                row_cells[2].text = self.format_date(action.get('due_date', ''))
                row_cells[3].text = action.get('status', 'Open')
            
            doc.add_paragraph()  # Add spacing
        
        # Decisions
        if meeting.get('decisions'):
            doc.add_heading('Decisions Made', level=1)
            
            for i, decision in enumerate(meeting['decisions'], 1):
                decision_para = doc.add_paragraph()
                decision_para.add_run(f"Decision {i}: ").bold = True
                decision_para.add_run(decision.get('decision', ''))
                
                if decision.get('impact'):
                    impact_para = doc.add_paragraph()
                    impact_para.add_run("Impact: ").bold = True
                    impact_para.add_run(decision.get('impact'))
                
                doc.add_paragraph()  # Add spacing between decisions
        
        # Notes
        if meeting.get('notes'):
            doc.add_heading('Additional Notes', level=1)
            notes_clean = self.clean_html_content(meeting['notes'])
            doc.add_paragraph(notes_clean)
        
        # Next Meeting
        next_meeting = meeting.get('next_meeting', {})
        if next_meeting:
            doc.add_heading('Next Meeting', level=1)
            
            next_para = doc.add_paragraph()
            next_para.add_run("Date: ").bold = True
            next_para.add_run(self.format_datetime(next_meeting.get('date', '')))
            
            if next_meeting.get('location'):
                location_para = doc.add_paragraph()
                location_para.add_run("Location: ").bold = True
                location_para.add_run(next_meeting.get('location'))
            
            if next_meeting.get('agenda_preview'):
                agenda_para = doc.add_paragraph()
                agenda_para.add_run("Agenda Items: ").bold = True
                agenda_para.add_run(next_meeting.get('agenda_preview'))
        
        # Footer
        doc.add_paragraph()
        footer_para = doc.add_paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.runs[0].font.size = 8
        
        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer


def get_meeting_filename(meeting: Dict[str, Any], format_type: str) -> str:
    """Generate a clean filename for the meeting export"""
    title = meeting.get('title', 'Meeting')
    # Clean the title for use in filename
    import re
    clean_title = re.sub(r'[^\w\s-]', '', title).strip()
    clean_title = re.sub(r'[-\s]+', '-', clean_title)
    
    date_str = ""
    if meeting.get('date'):
        try:
            dt = datetime.fromisoformat(meeting['date'].replace('Z', '+00:00'))
            date_str = dt.strftime('%Y%m%d')
        except:
            pass
    
    filename = f"meeting-{clean_title}"
    if date_str:
        filename += f"-{date_str}"
    
    filename += f".{format_type}"
    return filename